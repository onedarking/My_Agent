"""Agent tools - all file processing operations"""

import os
import re
import shutil
import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
import pdfplumber
from docx import Document

from config import RAW_DIR, PROCESSED_DIR, LOGS_DIR


# ─── Logging ──────────────────────────────────────────────

def log_action(action: str, file: str, status: str, detail: str = ""):
    """Append to process log."""
    log_file = LOGS_DIR / "process_log.csv"
    entry = {
        "timestamp": datetime.now().isoformat(),
        "action": action,
        "file": file,
        "status": status,
        "detail": detail,
    }
    if log_file.exists():
        df = pd.read_csv(log_file)
        df = pd.concat([df, pd.DataFrame([entry])], ignore_index=True)
    else:
        df = pd.DataFrame([entry])
    df.to_csv(log_file, index=False)
    return entry


def get_logs(limit: int = 50) -> list:
    log_file = LOGS_DIR / "process_log.csv"
    if not log_file.exists():
        return []
    df = pd.read_csv(log_file)
    return df.tail(limit).to_dict("records")


# ─── File processing ─────────────────────────────────────

def ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)


def pdf_to_word(pdf_path: str, output_path: Optional[str] = None) -> dict:
    """Convert PDF to Word document."""
    pdf_path = Path(pdf_path)
    if not output_path:
        output_path = str(PROCESSED_DIR / f"{pdf_path.stem}.docx")

    doc = Document()
    # Set default font for Chinese support
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = 110000  # 11pt in half-points

    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                # Extract text with better params for Chinese
                text = page.extract_text(
                    x_tolerance=3,      # Reduce jitter tolerance for CJK
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=True, # Better for multi-column
                ) or ""
                if text.strip():
                    p = doc.add_paragraph(text.strip())
                    # Set paragraph font
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = 110000

                # Extract tables with merged cells handling
                tables = page.extract_tables()
                for table in tables:
                    if not table or len(table) < 1:
                        continue
                    rows_count = len(table)
                    cols_count = max(len(r) for r in table) if table else 0
                    if rows_count > 0 and cols_count > 0:
                        doc_table = doc.add_table(rows=rows_count, cols=cols_count)
                        doc_table.style = 'Table Grid'
                        for i, row in enumerate(table):
                            for j, cell in enumerate(row):
                                if j < cols_count:
                                    cell_text = str(cell or "").strip()
                                    doc_table.cell(i, j).text = cell_text
                        doc.add_paragraph()

        doc.save(str(output_path))
        log_action("pdf_to_word", pdf_path.name, "success", str(output_path))
        return {"status": "success", "output": str(output_path)}
    except Exception as e:
        log_action("pdf_to_word", pdf_path.name, "failed", str(e))
        return {"status": "failed", "error": str(e)}


def pdf_to_excel(pdf_path: str, output_path: Optional[str] = None) -> dict:
    """Extract tables from PDF to Excel."""
    pdf_path = Path(pdf_path)
    if not output_path:
        output_path = str(PROCESSED_DIR / f"{pdf_path.stem}.xlsx")

    all_tables = []
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                for j, table in enumerate(tables):
                    if table and len(table) > 1:
                        headers = table[0]
                        data = table[1:]
                        df = pd.DataFrame(data, columns=headers)
                        df["来源页"] = i + 1
                        all_tables.append(df)

        if all_tables:
            merged = pd.concat(all_tables, ignore_index=True)
            merged.to_excel(output_path, index=False)
            log_action("pdf_to_excel", pdf_path.name, "success", str(output_path))
            return {"status": "success", "rows": len(merged), "output": str(output_path)}
        else:
            # If no tables found, split text by line -> rows
            with pdfplumber.open(str(pdf_path)) as pdf:
                lines = []
                for page in pdf.pages:
                    text = page.extract_text(
                        x_tolerance=3, y_tolerance=3,
                        keep_blank_chars=False, use_text_flow=True,
                    ) or ""
                    for line in text.split("\n"):
                        line = line.strip()
                        if line:
                            lines.append(line)
            df = pd.DataFrame({"行号": range(1, len(lines)+1), "内容": lines})
            df.to_excel(output_path, index=False)
            log_action("pdf_to_excel", pdf_path.name, "success", str(output_path))
            return {"status": "success", "rows": 1, "output": str(output_path)}
    except Exception as e:
        log_action("pdf_to_excel", pdf_path.name, "failed", str(e))
        return {"status": "failed", "error": str(e)}


def extract_report_info(pdf_path: str) -> dict:
    """Use LLM to extract structured info from a report PDF."""
    pdf_path = Path(pdf_path)
    with pdfplumber.open(str(pdf_path)) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages[:10])

    log_action("extract_info", pdf_path.name, "processing", f"extracted {len(text)} chars")
    return {"status": "llm_ready", "text_preview": text[:2000], "full_length": len(text)}


def clean_excel_data(excel_path: str, output_path: Optional[str] = None) -> dict:
    """Clean Excel: remove dupes, fillna, basic outlier detection."""
    excel_path = Path(excel_path)
    if not output_path:
        out_name = excel_path.stem  # Without extension
        output_path = str(PROCESSED_DIR / f"cleaned_{out_name}.xlsx")

    try:
        ext = str(excel_path).lower()
        if ext.endswith('.csv'):
            df = pd.read_csv(str(excel_path), engine='python')
        else:
            df = pd.read_excel(str(excel_path))
        report = {"original_rows": len(df), "original_cols": len(df.columns)}

        # Remove duplicates
        before = len(df)
        df = df.drop_duplicates()
        report["duplicates_removed"] = before - len(df)

        # Fill numeric NaN with median
        num_cols = df.select_dtypes(include="number").columns
        for col in num_cols:
            df[col] = df[col].fillna(df[col].median())

        # Fill text NaN
        text_cols = df.select_dtypes(include="object").columns
        for col in text_cols:
            df[col] = df[col].fillna("")

        report["cleaned_rows"] = len(df)

        # Outlier detection (basic IQR)
        outliers = {}
        for col in num_cols:
            q1 = df[col].quantile(0.25)
            q3 = df[col].quantile(0.75)
            iqr = q3 - q1
            outlier_mask = (df[col] < q1 - 1.5 * iqr) | (df[col] > q3 + 1.5 * iqr)
            if outlier_mask.sum() > 0:
                outliers[col] = int(outlier_mask.sum())
        report["outlier_columns"] = outliers

        df.to_excel(output_path, index=False)
        log_action("excel_clean", excel_path.name, "success", str(output_path))
        report["status"] = "success"
        report["output"] = str(output_path)
        return report
    except Exception as e:
        log_action("excel_clean", excel_path.name, "failed", str(e))
        return {"status": "failed", "error": str(e)}


def batch_rename(directory: str, pattern: str, prefix: str = "") -> list:
    """Batch rename files matching pattern. e.g. pattern='*.pdf', prefix='report_'"""
    dir_path = Path(directory)
    results = []
    for i, f in enumerate(sorted(dir_path.glob(pattern)), 1):
        ext = f.suffix
        new_name = f"{prefix}{i:03d}{ext}"
        new_path = f.with_name(new_name)
        f.rename(new_path)
        results.append({"old": f.name, "new": new_name})
        log_action("rename", f.name, "success", new_name)
    return results


def organize_by_type(directory: str) -> dict:
    """Organize files into subdirectories by extension type."""
    dir_path = Path(directory)
    categories = {
        "documents": [".pdf", ".docx", ".doc", ".txt", ".md"],
        "spreadsheets": [".xlsx", ".xls", ".csv"],
        "images": [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".svg"],
        "archives": [".zip", ".rar", ".7z", ".tar", ".gz"],
        "code": [".py", ".js", ".html", ".css", ".json", ".yaml", ".yml"],
    }

    moved = {}
    for f in dir_path.iterdir():
        if f.is_file():
            ext = f.suffix.lower()
            target = None
            for cat, exts in categories.items():
                if ext in exts:
                    target = cat
                    break
            if not target:
                target = "other"

            target_dir = dir_path / target
            ensure_dir(target_dir)
            dest = target_dir / f.name
            # Handle name conflicts
            if dest.exists():
                stem = dest.stem
                dest = target_dir / f"{stem}_{uuid.uuid4().hex[:4]}{dest.suffix}"

            shutil.move(str(f), str(dest))
            moved.setdefault(target, []).append(f.name)
            log_action("organize", f.name, "success", str(target_dir))

    return {k: len(v) for k, v in moved.items()}


def get_data_summary() -> dict:
    """Get summary of all data for dashboard."""
    summary = {
        "raw_files": len(list(RAW_DIR.iterdir())) if RAW_DIR.exists() else 0,
        "processed_files": len(list(PROCESSED_DIR.iterdir())) if PROCESSED_DIR.exists() else 0,
        "log_entries": len(get_logs(9999)),
        "raw_dir": str(RAW_DIR),
        "processed_dir": str(PROCESSED_DIR),
    }
    # File types breakdown
    types = {}
    for f in RAW_DIR.iterdir():
        if f.is_file():
            ext = f.suffix.lower() or "no_ext"
            types[ext] = types.get(ext, 0) + 1
    summary["file_types"] = types
    return summary
